"""
K3 MRL Project Indexer — EmbeddingGemma-300M Edition
Builds mrl_index.pkl using sentence-transformers locally.
Zero API costs. Zero data leaves the machine. Sub-15ms/embedding.

Sources: Night-2 DOCX, Night-3 MD, brain artifacts, KI artifacts, skills, deep research.

EmbeddingGemma-300M: 768-dim with native MRL truncation support.
Uses prompt_name="Retrieval-document" for optimal retrieval quality.
"""
import os
import glob
import hashlib
import pickle
import time
import re
import numpy as np
from pathlib import Path

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("[WARN] python-docx not installed — skipping DOCX files")

# ============================================================
# CONFIG
# ============================================================
MODEL_NAME = "google/embeddinggemma-300M"
TARGET_DIMENSION = 768  # Native dim; can truncate to 256/128 via MRL
BATCH_SIZE = 64         # sentence-transformers handles batching internally
SAVE_INTERVAL = 200

SRC_DIR = r"C:\K3_Firehose\k3-mcp-toolbox-public\k3-mcp-toolbox\src"
OUT_FILE = os.path.join(SRC_DIR, "mrl_index.pkl")

# Source directories
DOCX_DIR = r"C:\K3_Firehose\realtor-research\research night 2"
NIGHT3_MD_DIR = r"C:\K3_Firehose\last_night_research\3-20-26 research\md"
BRAIN_DIR = r"C:\Users\fandr\.gemini\antigravity\brain\21d5787a-9886-4798-99d0-046410654cae"
KI_DIR = r"C:\Users\fandr\.gemini\antigravity\knowledge"
SKILLS_DIR = r"C:\K3_Firehose\.agent\skills"
DEEP_RESEARCH_DIR = r"C:\K3_Firehose\90 deep research\md"

# ============================================================
# HELPERS
# ============================================================

def extract_docx_text(path):
    if not HAS_DOCX:
        return ""
    doc = docx.Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def sanitize(text):
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    text = re.sub(r"\s+", " ", text).strip()
    return text


def md5(text):
    return hashlib.md5(text.encode("utf-8")).hexdigest()


def smart_chunk(text, limit=2000):
    """Chunk text into pieces under `limit` chars (~500 tokens).
    EmbeddingGemma handles up to 8192 tokens, so we can be generous."""
    if len(text) <= limit:
        return [(text, "::main")]
    chunks, current, current_len, idx = [], [], 0, 0
    for line in text.split("\n"):
        ll = len(line) + 1
        if current_len + ll > limit and current:
            chunks.append(("\n".join(current), f"::chunk[{idx}]"))
            idx += 1
            current, current_len = [], 0
        current.append(line)
        current_len += ll
    if current:
        chunks.append(("\n".join(current), f"::chunk[{idx}]"))
    return chunks


# ============================================================
# MAIN
# ============================================================
def main():
    print("=" * 60)
    print("K3 MRL Index Builder — EmbeddingGemma-300M Edition")
    print("=" * 60)
    print(f"[CONFIG] Model: {MODEL_NAME}")
    print(f"[CONFIG] Dimension: {TARGET_DIMENSION} (MRL-native)")
    print(f"[CONFIG] Output: {OUT_FILE}")
    print()

    # Load model
    print("[LOADING] EmbeddingGemma-300M via sentence-transformers...")
    t0 = time.time()
    try:
        from sentence_transformers import SentenceTransformer
        model = SentenceTransformer(MODEL_NAME)
        print(f"[MODEL OK] Loaded in {time.time()-t0:.1f}s (<200MB RAM)")
    except Exception as e:
        print(f"[FATAL] Model load failed: {e}")
        print("        Run: pip install sentence-transformers")
        return

    # Quick test
    test_vec = model.encode("connectivity test", prompt_name="Retrieval-document",
                            truncate_dim=TARGET_DIMENSION, normalize_embeddings=True)
    print(f"[TEST OK] Output shape: {test_vec.shape}, norm: {np.linalg.norm(test_vec):.4f}")

    # Load existing index
    if os.path.exists(OUT_FILE):
        with open(OUT_FILE, "rb") as f:
            index = pickle.load(f)
        print(f"[LOADED] Existing index with {len(index)} entries")
    else:
        index = {}
        print("[FRESH BUILD] No existing index found")

    # --- Collect documents ---
    documents = []

    # 1. Night-2 DOCX
    if HAS_DOCX and os.path.isdir(DOCX_DIR):
        docx_files = glob.glob(os.path.join(DOCX_DIR, "*.docx"))
        for df in docx_files:
            try:
                text = extract_docx_text(df)
                fname = os.path.basename(df)
                documents.append({"id": fname, "text": f"Source: {fname}\n\n{text}", "type": "research"})
            except Exception as e:
                print(f"  [WARN] DOCX error {os.path.basename(df)}: {e}")
        print(f"[SOURCES] {len(docx_files)} DOCX research files (Night-2)")
    else:
        print("[SOURCES] 0 DOCX (skipped)")

    # 2. Night-3 MD research
    night3_count = 0
    if os.path.isdir(NIGHT3_MD_DIR):
        for mf in glob.glob(os.path.join(NIGHT3_MD_DIR, "*.md")):
            try:
                text = Path(mf).read_text(encoding="utf-8", errors="ignore")
                fname = os.path.basename(mf)
                documents.append({"id": f"night3:{fname}", "text": f"Research (3-20-26): {fname}\n\n{text}", "type": "research"})
                night3_count += 1
            except Exception as e:
                print(f"  [WARN] Night-3 MD error {os.path.basename(mf)}: {e}")
    print(f"[SOURCES] {night3_count} MD research files (Night-3)")

    # 3. Brain artifacts
    brain_count = 0
    if os.path.isdir(BRAIN_DIR):
        md_files = glob.glob(os.path.join(BRAIN_DIR, "*.md"))
        for mf in md_files:
            try:
                text = Path(mf).read_text(encoding="utf-8", errors="ignore")
                fname = os.path.basename(mf)
                documents.append({"id": f"brain:{fname}", "text": f"Artifact: {fname}\n\n{text}", "type": "brain"})
                brain_count += 1
            except Exception as e:
                print(f"  [WARN] Brain error {os.path.basename(mf)}: {e}")
    print(f"[SOURCES] {brain_count} brain artifacts")

    # 4. KI artifacts
    ki_count = 0
    for ki_folder in Path(KI_DIR).iterdir():
        if not ki_folder.is_dir():
            continue
        artifacts_dir = ki_folder / "artifacts"
        if artifacts_dir.exists():
            for af in artifacts_dir.rglob("*.md"):
                try:
                    text = af.read_text(encoding="utf-8", errors="ignore")
                    rel = af.relative_to(ki_folder)
                    fname = f"KI:{ki_folder.name}/{rel}"
                    documents.append({"id": fname, "text": f"Knowledge: {fname}\n\n{text}", "type": "ki"})
                    ki_count += 1
                except Exception:
                    pass
    print(f"[SOURCES] {ki_count} knowledge item artifacts")

    # 4b. Deep Research
    dr_count = 0
    if os.path.isdir(DEEP_RESEARCH_DIR):
        for mf in glob.glob(os.path.join(DEEP_RESEARCH_DIR, "*.md")):
            try:
                text = Path(mf).read_text(encoding="utf-8", errors="ignore")
                fname = os.path.basename(mf)
                documents.append({"id": f"deep-research:{fname}", "text": f"Deep Research: {fname}\n\n{text}", "type": "deep-research"})
                dr_count += 1
            except Exception as e:
                print(f"  [WARN] Deep Research error {os.path.basename(mf)}: {e}")
    print(f"[SOURCES] {dr_count} deep research files")

    # 5. Skills (SKILL.md files)
    skill_count = 0
    for skill_dir in Path(SKILLS_DIR).iterdir():
        if not skill_dir.is_dir():
            continue
        skill_md = skill_dir / "SKILL.md"
        if skill_md.exists():
            try:
                text = skill_md.read_text(encoding="utf-8", errors="ignore")
                fname = f"skill:{skill_dir.name}"
                documents.append({"id": fname, "text": f"Skill: {skill_dir.name}\n\n{text}", "type": "skill"})
                skill_count += 1
            except Exception:
                pass
    print(f"[SOURCES] {skill_count} skills")
    print(f"[TOTAL]   {len(documents)} documents to process")

    # --- Chunk and diff against existing index ---
    pending_paths = []
    pending_texts = []
    pending_hashes = []

    for doc in documents:
        for chunk_text, suffix in smart_chunk(doc["text"]):
            clean = sanitize(chunk_text)
            if not clean or len(clean) < 20:
                continue
            full_path = f"{doc['id']}{suffix}"
            new_hash = md5(clean)
            if full_path in index and index[full_path].get("hash") == new_hash:
                continue
            pending_paths.append(full_path)
            pending_texts.append(clean)
            pending_hashes.append(new_hash)

    print(f"[CHUNKS]  {len(pending_texts)} need embedding ({len(index)} cached)")

    if not pending_texts:
        print("[DONE] Index is up to date!")
        return

    # --- Batch embed with EmbeddingGemma ---
    print(f"[EMBEDDING] Processing {len(pending_texts)} chunks in batches of {BATCH_SIZE}...")
    start_time = time.time()
    total = 0
    unsaved = 0

    for bi in range(0, len(pending_texts), BATCH_SIZE):
        batch_texts = pending_texts[bi:bi + BATCH_SIZE]
        batch_paths = pending_paths[bi:bi + BATCH_SIZE]
        batch_hashes = pending_hashes[bi:bi + BATCH_SIZE]
        batch_num = bi // BATCH_SIZE + 1
        n_batches = (len(pending_texts) + BATCH_SIZE - 1) // BATCH_SIZE

        # Encode batch — sentence-transformers handles internal batching
        vectors = model.encode(
            batch_texts,
            prompt_name="Retrieval-document",
            truncate_dim=TARGET_DIMENSION,
            normalize_embeddings=True,
            show_progress_bar=False
        )

        for path, h, vec in zip(batch_paths, batch_hashes, vectors):
            index[path] = {
                "vector": vec.astype(np.float32),
                "hash": h,
                "snippet": pending_texts[pending_paths.index(path)][:200]
            }
            total += 1
            unsaved += 1

        elapsed = time.time() - start_time
        rate = total / elapsed if elapsed > 0 else 0
        remaining = (len(pending_texts) - bi - len(batch_texts)) / rate if rate > 0 else 0
        print(f"  [{batch_num}/{n_batches}] +{len(batch_texts)} ({total} done, {rate:.1f}/s, ~{remaining:.0f}s left)")

        if unsaved >= SAVE_INTERVAL:
            _save(index)
            unsaved = 0

    _save(index)
    elapsed = time.time() - start_time

    print(f"\n{'='*60}")
    print(f"[COMPLETE] mrl_index.pkl — EmbeddingGemma-300M Edition")
    print(f"  Entries:   {len(index)}")
    print(f"  New:       {total}")
    print(f"  Time:      {elapsed:.1f}s ({elapsed/60:.1f} min)")
    print(f"  Rate:      {total/elapsed:.1f} embeddings/sec")
    print(f"  Path:      {OUT_FILE}")

    # Source breakdown
    types = {}
    for k in index:
        if k.startswith("night3:"):
            t = "night3-research"
        elif k.startswith("deep-research:"):
            t = "deep-research"
        elif k.startswith("brain:"):
            t = "brain"
        elif k.startswith("KI:"):
            t = "ki"
        elif k.startswith("skill:"):
            t = "skill"
        else:
            t = "night2-research"
        types[t] = types.get(t, 0) + 1
    print(f"\n  Source breakdown:")
    for t, c in sorted(types.items()):
        print(f"    {t}: {c} chunks")


def _save(index):
    tmp = OUT_FILE + ".tmp"
    with open(tmp, "wb") as f:
        pickle.dump(index, f, protocol=pickle.HIGHEST_PROTOCOL)
    if os.path.exists(OUT_FILE):
        os.remove(OUT_FILE)
    os.rename(tmp, OUT_FILE)
    print(f"  [SAVED] {len(index)} entries")


if __name__ == "__main__":
    main()
